import os
import io
import json
import re
import uuid
from typing import List, Dict, Any

import streamlit as st
from pydantic import BaseModel, Field

# --- NLP / IR ---
import numpy as np
from sentence_transformers import SentenceTransformer
import faiss
from rank_bm25 import BM25Okapi
from rapidfuzz import fuzz
from unidecode import unidecode

# --- PDF / DOCX / PDF Export ---
import pdfplumber
from pypdf import PdfReader
from docx import Document
from docx.shared import Pt
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4

# --- HTTP / LLM ---
import httpx

# =========================
# Config & Secrets
# =========================
st.set_page_config(page_title="DR.M — Petições com RAG", layout="wide")

# Provider & Models
DEFAULT_PROVIDER = st.secrets.get("LLM_PROVIDER", os.getenv("LLM_PROVIDER", "sabia")).lower()
OPENAI_API_KEY = st.secrets.get("OPENAI_API_KEY", os.getenv("OPENAI_API_KEY", ""))
SABIA_API_KEY = st.secrets.get("SABIA_API_KEY", os.getenv("SABIA_API_KEY", ""))
SABIA_BASE_URL = st.secrets.get("SABIA_BASE_URL", os.getenv("SABIA_BASE_URL", ""))

MODEL_GPT = st.secrets.get("MODEL_GPT", os.getenv("MODEL_GPT", "gpt-4o-mini"))
MODEL_SABIA = st.secrets.get("MODEL_SABIA", os.getenv("MODEL_SABIA", "sabia-3.1"))

EMBEDDING_MODEL = st.secrets.get("EMBEDDING_MODEL", os.getenv("EMBEDDING_MODEL", "sentence-transformers/all-MiniLM-L6-v2"))

# RAG params
CHUNK_SIZE = int(st.secrets.get("CHUNK_SIZE", os.getenv("CHUNK_SIZE", 1100)))
CHUNK_OVERLAP = int(st.secrets.get("CHUNK_OVERLAP", os.getenv("CHUNK_OVERLAP", 200)))
TOP_K = int(st.secrets.get("TOP_K", os.getenv("TOP_K", 6)))
USE_BM25 = (str(st.secrets.get("USE_BM25", os.getenv("USE_BM25", "true"))).lower() == "true")

TEMP_WRITER = float(st.secrets.get("TEMP_WRITER", os.getenv("TEMP_WRITER", 0.4)))
TEMP_REVIEW_MERIT = float(st.secrets.get("TEMP_REVIEW_MERIT", os.getenv("TEMP_REVIEW_MERIT", 0.2)))
TEMP_REVIEW_PROC = float(st.secrets.get("TEMP_REVIEW_PROC", os.getenv("TEMP_REVIEW_PROC", 0.1)))
TEMP_REVIEW_FORMAT = float(st.secrets.get("TEMP_REVIEW_FORMAT", os.getenv("TEMP_REVIEW_FORMAT", 0.2)))

# =========================
# Models de dados
# =========================
class Party(BaseModel):
    nome: str
    tipo: str = Field(description="Pessoa Física ou Jurídica")
    cpf_cnpj: str | None = None
    endereco: str | None = None
    email: str | None = None

class CaseInfo(BaseModel):
    foro_comarca: str
    area: str = "Cível"
    autor: Party
    reu: Party
    causa_de_pedir: str
    pedidos: List[str]
    valor_da_causa: str | None = None
    urgencia: bool = False
    observacoes: str | None = None

# =========================
# Utils
# =========================
def normalize_text(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").replace("\x00", "")).strip()

def chunk_text(text: str, size: int, overlap: int) -> List[str]:
    words = text.split()
    chunks, start = [], 0
    while start < len(words):
        end = min(len(words), start + size)
        chunks.append(" ".join(words[start:end]))
        if end == len(words): break
        start = max(0, end - overlap)
    return chunks

def read_pdf_text(file_bytes: bytes) -> str:
    # 1) pypdf
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        txt = "\n".join((page.extract_text() or "") for page in reader.pages)
        txt = normalize_text(txt)
        if txt: return txt
    except Exception:
        pass
    # 2) pdfplumber
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            txt = "\n".join((pg.extract_text() or "") for pg in pdf.pages)
        txt = normalize_text(txt)
        if txt: return txt
    except Exception:
        pass
    # 3) Sem OCR no Cloud por padrão
    return ""

# =========================
# VectorIndex (FAISS + BM25)
# =========================
class VectorIndex:
    def __init__(self, embedder_name: str):
        self.model = SentenceTransformer(embedder_name)
        self.index = None
        self.doc_meta: List[Dict[str, Any]] = []
        self.emb_dim = self.model.get_sentence_embedding_dimension()
        self.bm25 = None

    def _build_bm25(self):
        tokenized = [m["text"].lower().split() for m in self.doc_meta]
        self.bm25 = BM25Okapi(tokenized)

    def add(self, texts: List[str], metas: List[Dict[str, Any]]):
        embs = self.model.encode(texts, convert_to_numpy=True, normalize_embeddings=True).astype(np.float32)
        if self.index is None:
            self.index = faiss.IndexFlatIP(self.emb_dim)
        self.index.add(embs)
        self.doc_meta.extend(metas)
        if USE_BM25:
            self._build_bm25()
        # reindexa idx
        for i, m in enumerate(self.doc_meta):
            m["idx"] = i

    def search(self, query: str, top_k: int) -> List[Dict[str, Any]]:
        if not self.doc_meta or self.index is None:
            return []
        q = self.model.encode([query], convert_to_numpy=True, normalize_embeddings=True).astype(np.float32)
        D, I = self.index.search(q, min(top_k, len(self.doc_meta)))
        vec_res = []
        for score, idx in zip(D[0].tolist(), I[0].tolist()):
            if idx < 0: continue
            meta = dict(self.doc_meta[idx]); meta["score_vec"] = float(score)
            vec_res.append(meta)

        if USE_BM25 and self.bm25 is not None:
            bm = self.bm25.get_scores(query.lower().split())
            bmin, bmax = float(np.min(bm)), float(np.max(bm))
            denom = (bmax - bmin) or 1.0
            bm_norm = (bm - bmin) / denom
            fused = {}
            for r in vec_res:
                idx = r["idx"]
                r["score_fused"] = 0.5 * r["score_vec"] + 0.5 * float(bm_norm[idx])
                fused[idx] = r
            for idx in np.argsort(-bm_norm)[:top_k].tolist():
                if idx not in fused and idx < len(self.doc_meta):
                    r = dict(self.doc_meta[idx])
                    r["score_vec"] = 0.0
                    r["score_fused"] = float(bm_norm[idx])
                    fused[idx] = r
            return sorted(fused.values(), key=lambda x: x["score_fused"], reverse=True)[:top_k]
        return vec_res[:top_k]

# =========================
# LLM Providers
# =========================
class LLMMessage(BaseModel):
    role: str
    content: str

def sabia_chat(messages: List[LLMMessage], model: str, temperature: float, max_tokens: int) -> str:
    base = (SABIA_BASE_URL or "").strip()
    key = (SABIA_API_KEY or "").strip()
    if not base.startswith("http"):
        base = "https://" + base
    if not base or not key:
        raise RuntimeError("Configure SABIA_BASE_URL e SABIA_API_KEY em st.secrets ou variáveis de ambiente.")
    headers = {"Authorization": f"Bearer {key}", "Content-Type": "application/json"}
    payload = {
        "model": model,
        "messages": [m.model_dump() for m in messages],
        "temperature": temperature,
        "max_tokens": max_tokens
    }
    with httpx.Client(timeout=120) as c:
        r = c.post(base, headers=headers, json=payload)
        r.raise_for_status()
        data = r.json()
    try:
        return data["choices"][0]["message"]["content"]
    except Exception:
        return json.dumps(data, ensure_ascii=False)

def openai_chat(messages: List[LLMMessage], model: str, temperature: float, max_tokens: int) -> str:
    # sem o SDK — usando REST via httpx para evitar dependência extra
    key = (OPENAI_API_KEY or "").strip()
    if not key:
        raise RuntimeError("Configure OPENAI_API_KEY em st.secrets ou variável de ambiente.")
    headers = {"Authorization": f"Bearer {key}", "Content-Type": "application/json"}
    payload = {
        "model": model,
        "messages": [m.model_dump() for m in messages],
        "temperature": temperature,
        "max_tokens": max_tokens
    }
    url = "https://api.openai.com/v1/chat/completions"
    with httpx.Client(timeout=120) as c:
        r = c.post(url, headers=headers, json=payload)
        r.raise_for_status()
        data = r.json()
    return data["choices"][0]["message"]["content"]

def llm_chat(provider: str, messages: List[LLMMessage], temperature: float, max_tokens: int) -> str:
    if provider == "gpt":
        return openai_chat(messages, MODEL_GPT, temperature, max_tokens)
    return sabia_chat(messages, MODEL_SABIA, temperature, max_tokens)

# =========================
# Prompts
# =========================
PROMPT_WRITER = """Você é um advogado sênior brasileiro. Redija uma PETIÇÃO completa e objetiva, com as seções:
1) Endereçamento
2) Qualificação das partes
3) Dos fatos
4) Do direito (fundamentos jurídicos)
5) Dos pedidos (numeração)
6) Valor da causa
7) Requerimentos finais (citação/justiça gratuita/tutela/antecipação, se aplicável)
8) Rol de documentos anexos

Regras:
- Linguagem técnica clara e respeitosa.
- Utilize as **fontes recuperadas** quando pertinentes e cite-as no formato [Fonte: <doc> <chunk_id>] ao final do parágrafo correspondente.
- Não invente fatos. Baseie-se apenas no resumo do caso, nos pedidos e nos trechos recuperados.
- Se houver urgência, inclua seção de tutela de urgência.
"""

PROMPT_REVIEW_MERIT = """Você é um revisor jurídico focado em MÉRITO. Avalie criticamente a minuta abaixo.
Responda em JSON com:
{
 "issues": ["..."],
 "suggested_fixes": ["..."],
 "quality_notes": {"clareza":0-10,"aderencia":0-10,"riscos":"..."}
}
Seja específico. Aponte fundamentos legais, súmulas e precedentes aplicáveis. Não reescreva a peça inteira.
"""

PROMPT_REVIEW_PROC = """Você é um revisor jurídico focado em PROCEDIMENTO/FORMALIDADES. Verifique:
- Endereçamento adequado (foro/competência)
- Qualificação das partes
- Estrutura (fatos, direito, pedidos)
- Pedidos compatíveis (citação, tutela, justiça gratuita, provas)
- Valor da causa
- Lista de anexos
Responda em JSON com:
{
 "issues": ["..."],
 "suggested_fixes": ["..."],
 "quality_notes": {"conformidade":0-10,"riscos_processuais":"..."}
}
"""

PROMPT_REVIEW_FORMAT = """Você é um revisor de FORMATAÇÃO. Ajuste a minuta para:
- Cabeçalhos/seções bem definidas
- Parágrafos com espaçamento adequado
- Numeração e alinhamento dos pedidos
- Preservar citações [Fonte: <doc> <chunk_id>]
- Português jurídico claro e padronizado
Responda com o **texto final formatado**, sem comentários.
"""

# =========================
# UI — Sidebar
# =========================
st.sidebar.title("Configurações")
provider = st.sidebar.selectbox("Provider", ["sabia", "gpt"], index=0 if DEFAULT_PROVIDER=="sabia" else 1)
st.sidebar.caption("Chaves devem estar em st.secrets ou variáveis de ambiente.")

with st.sidebar.expander("Parâmetros do Modelo", expanded=False):
    st.write("Temperaturas")
    t_writer = st.slider("Writer", 0.0, 1.0, float(TEMP_WRITER), 0.05)
    t_merit  = st.slider("Revisor Mérito", 0.0, 1.0, float(TEMP_REVIEW_MERIT), 0.05)
    t_proc   = st.slider("Revisor Proced.", 0.0, 1.0, float(TEMP_REVIEW_PROC), 0.05)
    t_format = st.slider("Revisor Formatação", 0.0, 1.0, float(TEMP_REVIEW_FORMAT), 0.05)
    top_k    = st.slider("TOP_K (RAG)", 1, 12, int(TOP_K), 1)

# =========================
# UI — Passo 1: Partes
# =========================
st.header("DR.M — Petições com RAG")
st.subheader("Passo 1 — Dados das Partes e Foro")

col1, col2 = st.columns(2)
with col1:
    autor_nome = st.text_input("Autor — Nome", value="Ricardo da Silva Santos")
    autor_tipo = st.selectbox("Autor — Tipo", ["Pessoa Física", "Pessoa Jurídica"], index=0)
    autor_doc  = st.text_input("Autor — CPF/CNPJ", value="000.000.000-00")
    autor_end  = st.text_input("Autor — Endereço", value="Endereço do Autor, Goiânia/GO")
    autor_email= st.text_input("Autor — E-mail", value="")
with col2:
    reu_nome = st.text_input("Réu — Nome", value="Operadora XYZ Saúde S.A.")
    reu_tipo = st.selectbox("Réu — Tipo", ["Pessoa Jurídica", "Pessoa Física"], index=0)
    reu_doc  = st.text_input("Réu — CNPJ/CPF", value="00.000.000/0001-00")
    reu_end  = st.text_input("Réu — Endereço", value="Sede da Operadora, São Paulo/SP")
    reu_email= st.text_input("Réu — E-mail", value="")

foro = st.text_input("Foro/Comarca", value="Foro Central de Goiânia/GO")
area = st.text_input("Área", value="Cível/Consumidor")

# =========================
# Passo 2: Motivo e Pedidos
# =========================
st.subheader("Passo 2 — Motivo/Resumo e Pedidos")
causa = st.text_area("Causa de Pedir (resumo objetivo do caso)", height=140,
                     value="Negativa injustificada de cobertura de tratamento prescrito por médico especialista; urgência comprovada por laudos.")
pedidos_raw = st.text_area("Pedidos (um por linha)", height=120,
                           value="Tutela de urgência para cobertura imediata\nCobertura integral do tratamento, exames e insumos\nDanos morais\nInversão do ônus da prova (art. 6º, VIII, CDC)\nCitação da ré")
valor_causa = st.text_input("Valor da Causa", value="R$ 50.000,00")
urgencia = st.checkbox("Há urgência (tutela)?", value=True)
obs = st.text_area("Observações (opcional)", height=80, value="")

# =========================
# Passo 3: Documentos
# =========================
st.subheader("Passo 3 — Upload de Documentos (PDFs)")
files = st.file_uploader("Envie jurisprudência, legislação, contratos ou peças anteriores (PDFs, multi-upload)", type=["pdf"], accept_multiple_files=True)
st.caption("Dica: PDFs com ‘texto selecionável’ melhoram o RAG. OCR não está habilitado no Cloud por padrão.")

# =========================
# Botão principal
# =========================
go = st.button("Gerar Petição")

# =========================
# Execução
# =========================
if go:
    with st.spinner("Processando..."):
        # Monta CaseInfo
        info = CaseInfo(
            foro_comarca=foro, area=area,
            autor=Party(nome=autor_nome, tipo=autor_tipo, cpf_cnpj=autor_doc, endereco=autor_end, email=autor_email or None),
            reu=Party(nome=reu_nome, tipo=reu_tipo, cpf_cnpj=reu_doc, endereco=reu_end, email=reu_email or None),
            causa_de_pedir=causa,
            pedidos=[normalize_text(p) for p in (pedidos_raw.splitlines() if pedidos_raw else []) if normalize_text(p)],
            valor_da_causa=valor_causa or None,
            urgencia=bool(urgencia),
            observacoes=obs or None
        )

        # Ingestão → Index
        vindex = VectorIndex(EMBEDDING_MODEL)
        all_passages: List[Dict[str, Any]] = []
        if files:
            for f in files:
                b = f.read()
                txt = read_pdf_text(b)
                if not txt:
                    st.warning(f"Sem texto extraído de: {f.name} (OCR não habilitado).")
                    continue
                chunks = chunk_text(txt, CHUNK_SIZE, CHUNK_OVERLAP)
                metas = []
                for i, ch in enumerate(chunks):
                    metas.append({
                        "doc_path": f.name,
                        "chunk_id": f"{f.name}#chunk={i}",
                        "text": ch,
                        "idx": len(vindex.doc_meta) + i
                    })
                vindex.add([m["text"] for m in metas], metas)
                all_passages.extend(metas)
        else:
            st.info("Nenhum PDF enviado — a peça será gerada sem citações a documentos anexos.")

        # Pesquisa (RAG)
        query = info.causa_de_pedir + " " + " ".join(info.pedidos)
        passages = vindex.search(query, top_k=top_k) if vindex.doc_meta else []

        # Writer
        resumo = f"""
FORO/COMARCA: {info.foro_comarca}
ÁREA: {info.area}

AUTOR: {info.autor.nome} ({info.autor.tipo})  CPF/CNPJ: {info.autor.cpf_cnpj or '-'}
ENDEREÇO AUTOR: {info.autor.endereco or '-'}

RÉU: {info.reu.nome} ({info.reu.tipo})  CPF/CNPJ: {info.reu.cpf_cnpj or '-'}
ENDEREÇO RÉU: {info.reu.endereco or '-'}

CAUSA DE PEDIR:
{info.causa_de_pedir}

PEDIDOS:
- """ + "\n- ".join(info.pedidos) + f"""

VALOR DA CAUSA: {info.valor_da_causa or '-'}

URGÊNCIA: {'SIM' if info.urgencia else 'NÃO'}
OBSERVAÇÕES: {info.observacoes or '-'}
"""
        passages_txt = "\n".join([f"- [{p['chunk_id']}] '{p['doc_path']}': {p['text'][:800]}..." for p in passages]) or "(Nenhum trecho recuperado)"
        draft = llm_chat(
            provider,
            [LLMMessage(role="system", content=PROMPT_WRITER),
             LLMMessage(role="user", content=f"RESUMO E DADOS DO CASO:\n{resumo}\n\nTRECHOS RECUPERADOS (use se pertinente):\n{passages_txt}")],
            temperature=t_writer, max_tokens=2200
        )

        # Reviewer — Mérito
        rev_merit_raw = llm_chat(
            provider,
            [LLMMessage(role="system", content=PROMPT_REVIEW_MERIT),
             LLMMessage(role="user", content=draft)],
            temperature=t_merit, max_tokens=1200
        )
        try:
            rev_merit = json.loads(rev_merit_raw)
        except Exception:
            rev_merit = {"issues": [], "suggested_fixes": [], "quality_notes": {"raw": rev_merit_raw[:800]}}

        # Reviewer — Procedimento
        rev_proc_raw = llm_chat(
            provider,
            [LLMMessage(role="system", content=PROMPT_REVIEW_PROC),
             LLMMessage(role="user", content=draft)],
            temperature=t_proc, max_tokens=1200
        )
        try:
            rev_proc = json.loads(rev_proc_raw)
        except Exception:
            rev_proc = {"issues": [], "suggested_fixes": [], "quality_notes": {"raw": rev_proc_raw[:800]}}

        # Reviewer — Formatação (3º revisor)
        payload_format = f"""
MINUTA:
{draft}

REVISOR MÉRITO (issues/sugestões):
{json.dumps(rev_merit, ensure_ascii=False)}

REVISOR PROCEDIMENTO (issues/sugestões):
{json.dumps(rev_proc, ensure_ascii=False)}
"""
        final_formatted = llm_chat(
            provider,
            [LLMMessage(role="system", content=PROMPT_REVIEW_FORMAT),
             LLMMessage(role="user", content=payload_format)],
            temperature=t_format, max_tokens=2200
        )

        st.success("Petições geradas! Veja abaixo.")

        # Exibir resultados
        with st.expander("Minuta (DRAFT)", expanded=True):
            st.write(draft)

        colA, colB = st.columns(2)
        with colA:
            st.subheader("Revisor — Mérito")
            st.json(rev_merit)
        with colB:
            st.subheader("Revisor — Procedimental")
            st.json(rev_proc)

        st.subheader("Versão Final (Formatada)")
        st.write(final_formatted)

        # Export DOCX
        def export_docx(text: str, passages: List[Dict[str, Any]], anexos: List[str]) -> bytes:
            doc = Document()
            style = doc.styles['Normal']; style.font.name = 'Calibri'; style.font.size = Pt(11)
            doc.add_heading("Petição — Versão Final", level=1)
            for para in text.split("\n\n"):
                doc.add_paragraph(para)

            if anexos:
                doc.add_heading("Rol de Documentos Anexos", level=2)
                for a in anexos:
                    doc.add_paragraph(f"- {os.path.basename(a)}")

            if passages:
                doc.add_heading("Apêndice — Trechos Recuperados (RAG)", level=2)
                for p in passages:
                    doc.add_paragraph(f"[{p['chunk_id']}] {os.path.basename(p['doc_path'])}")
                    doc.add_paragraph(p["text"][:1800] + ("..." if len(p["text"])>1800 else ""))

            mem = io.BytesIO()
            doc.save(mem)
            mem.seek(0)
            return mem.read()

        # Export PDF simples (texto corrido)
        def export_pdf(text: str) -> bytes:
            mem = io.BytesIO()
            c = canvas.Canvas(mem, pagesize=A4)
            width, height = A4
            x, y = 50, height - 50
            for line in text.split("\n"):
                if y < 60:
                    c.showPage()
                    y = height - 50
                c.drawString(x, y, line[:110])  # quebra simples
                y -= 14
            c.showPage()
            c.save()
            mem.seek(0)
            return mem.read()

        anexos_nomes = [f.name for f in files] if files else []
        docx_bytes = export_docx(final_formatted, passages, anexos_nomes)
        pdf_bytes  = export_pdf(final_formatted)

        colD, colE = st.columns(2)
        with colD:
            st.download_button("⬇️ Baixar DOCX", data=docx_bytes, file_name=f"Peticao_Final_{uuid.uuid4().hex[:8]}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        with colE:
            st.download_button("⬇️ Baixar PDF", data=pdf_bytes, file_name=f"Peticao_Final_{uuid.uuid4().hex[:8]}.pdf", mime="application/pdf")
